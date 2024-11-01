import gradio as gr
import requests
import json
import os
import io
import tempfile
import mimetypes
import time
import asyncio
import functools
import shutil
import random
import glob
from typing import List

from groq import Groq
from docx import Document
import fitz  # PyMuPDF để xử lý PDF
from PIL import Image, ImageDraw, ImageFont
from sentence_transformers import SentenceTransformer, util
from transformers import AutoTokenizer, AutoModel
from moviepy.editor import ImageClip, AudioFileClip, concatenate_videoclips
from edge_tts import Communicate  # Sử dụng Edge TTS
from huggingface_hub import login
from diffusers import DiffusionPipeline

# Cấu hình API keys
GROQ_API_KEY = os.environ.get("GROQ_API_KEY")
HF_API_KEY = os.environ.get("HF_API_KEY")

# Danh sách giọng đọc và ngôn ngữ
VOICES = ["vi-VN-HoaiMyNeural"]  # Thay đổi giọng đọc Tiếng Việt
LANGUAGES = ["Tiếng Anh", "Tiếng Việt"]
CONTENT_TYPES = ["podcast", "giới thiệu", "triết lý sống", "Phổ biến kiến thức thống kê"]

# Hướng dẫn mặc định cho từng loại nội dung
CONTENT_TYPE_INSTRUCTIONS = {
    "podcast": """
    Tone giọng: Gần gũi, thân thiện nhưng chuyên sâu, thể hiện sự am hiểu về chủ đề.
    Cấu trúc: 
    - Bắt đầu bằng một câu hỏi kích thích tư duy hoặc một câu chuyện mở màn gây tò mò.
    - Triển khai các luận điểm theo từng bước. Sử dụng câu từ mạnh mẽ, ví dụ điển hình hoặc những câu nói nổi tiếng.
    - Xây dựng các phần chuyển tiếp mượt mà giữa các ý.
    - Kết thúc podcast với một thông điệp sâu sắc, để lại sự suy ngẫm cho thính giả.
    Mục tiêu: Mang lại kiến thức giá trị, lôi cuốn thính giả tham gia suy nghĩ và cảm nhận sâu sắc về chủ đề.
    """,
    "giới thiệu": """
    Tone giọng: Chuyên nghiệp, gãy gọn nhưng vẫn có sự truyền cảm.
    Cấu trúc:
    - Bắt đầu với một câu khẳng định mạnh mẽ về đối tượng được giới thiệu.
    - Giải thích mục tiêu của phần giới thiệu, nhấn mạnh tầm quan trọng hoặc sự khác biệt.
    - Kết thúc với một lời kêu gọi hành động, khích lệ người nghe tiếp tục lắng nghe hoặc tham gia.
    Mục tiêu: Đưa ra thông tin cô đọng, hấp dẫn, khiến người nghe cảm thấy bị thu hút và muốn tìm hiểu thêm.
    """,
    "triết lý sống": """
    Tone giọng: Sâu sắc, truyền cảm hứng, mang tính chiêm nghiệm.
    Cấu trúc:
    - Bắt đầu bằng một câu hỏi sâu sắc hoặc ẩn dụ về cuộc sống.
    - Triển khai các luận điểm chặt chẽ, xen lẫn cảm xúc và những ví dụ đời thực hoặc những câu nói triết lý.
    - Kết thúc với một thông điệp sâu sắc, khơi dậy suy ngẫm cho người nghe.
    Mục tiêu: Khơi gợi suy nghĩ sâu sắc về cuộc sống, khiến người nghe tìm thấy ý nghĩa hoặc giá trị trong câu chuyện.
    """,
    "Phổ biến kiến thức Thống kê": """
    Tone giọng: Thân thiện, dễ hiểu, và mang tính giáo dục.
    Cấu trúc:
    - Bắt đầu với một câu hỏi hoặc một tình huống thực tế để thu hút sự chú ý.
    - Giải thích các khái niệm thống kê cơ bản một cách đơn giản và dễ hiểu, sử dụng ví dụ thực tế để minh họa.
    - Đưa ra các ứng dụng thực tế của thống kê trong đời sống hàng ngày hoặc trong các lĩnh vực cụ thể.
    - Kết thúc với một thông điệp khuyến khích người nghe áp dụng kiến thức thống kê vào cuộc sống.
    Mục tiêu: Giúp người nghe hiểu và yêu thích thống kê, thấy được giá trị và ứng dụng của nó trong cuộc sống.
    """
}

# Khởi tạo model sentence transformer
model = SentenceTransformer('all-MiniLM-L6-v2')

# Đăng nhập vào Hugging Face Hub
login(token=HF_API_KEY)

# Khởi tạo DiffusionPipeline
pipe = DiffusionPipeline.from_pretrained("black-forest-labs/FLUX.1-dev")
pipe.load_lora_weights("TDN-M/RetouchFLux")

@functools.lru_cache(maxsize=32)
def create_content(prompt, content_type, language):
    content_type_instructions = CONTENT_TYPE_INSTRUCTIONS.get(content_type, "")
    general_instructions = f"""
    Viết một kịch bản dựa trên các ý chính và ý tưởng sáng tạo từ yêu cầu của người dùng. Sử dụng giọng điệu trò chuyện và bao gồm bất kỳ bối cảnh hoặc giải thích cần thiết nào để làm cho nội dung dễ tiếp cận với khán giả.
    Bắt đầu kịch bản bằng cách nêu rõ chủ đề, tham chiếu đến tiêu đề hoặc đề mục trong văn bản đầu vào. Nếu văn bản đầu vào không có tiêu đề, hãy đưa ra một tóm tắt ngắn gọn về nội dung được đề cập để mở đầu.
    Bao gồm các định nghĩa và thuật ngữ rõ ràng, cùng với ví dụ cho tất cả các vấn đề chính.
    Không bao gồm bất kỳ placeholder nào trong ngoặc vuông như [Host] hoặc [Guest]. Thiết kế đầu ra của bạn để được đọc to - nó sẽ được chuyển đổi trực tiếp thành âm thanh.
    Chỉ có một người nói, bạn. Giữ đúng chủ đề và duy trì một luồng hấp dẫn. Văn phong triển khai một cách tự nhiên, .
    Kịch bản nên có khoảng 1000 từ, chia thành 5 đoạn rõ ràng. Hãy tuân theo những hướng dẫn cụ thể sau cho thể loại {content_type}:
    {content_type_instructions}
    Ngôn ngữ sử dụng: {language}
    """

    try:
        client = Groq(api_key=GROQ_API_KEY)
        retries = 3
        for attempt in range(retries):
            try:
                chat_completion = client.chat.completions.create(
                    model="mixtral-8x7b-32768",
                    messages=[
                        {"role": "system", "content": general_instructions},
                        {"role": "user", "content": prompt}
                    ],
                    temperature=0.7,
                    max_tokens=8000
                )
                return chat_completion.choices[0].message.content
            except Exception as e:
                if attempt < retries - 1:
                    time.sleep(2 ** attempt)  # Exponential backoff
                else:
                    return f"Lỗi khi tạo nội dung: {str(e)}"
    except Exception as e:
        return f"Lỗi khi tạo nội dung: {str(e)}"

def text_to_speech(text, voice, language):
    try:
        communicate = Communicate(text, voice)
        with tempfile.NamedTemporaryFile(delete=False, suffix=".mp3") as temp_audio_file:
            communicate.save(temp_audio_file.name)
        return temp_audio_file.name
    except Exception as e:
        return f"Lỗi khi chuyển đổi văn bản thành giọng nói: {str(e)}"

def create_docx(content, output_path):
    doc = Document()
    doc.add_paragraph(content)
    doc.save(output_path)

def process_pdf(file_path):
    doc = fitz.open(file_path)
    text = ""
    for page in doc:
        text += page.get_text()
    return text

def process_docx(file_path):
    doc = Document(file_path)
    text = ""
    for para in doc.paragraphs:
        text += para.text
    return text

def split_into_paragraphs(script_content):
    try:
        client = Groq(api_key=GROQ_API_KEY)
        response = client.chat.completions.create(
            model="mixtral-8x7b-32768",
            messages=[
                {"role": "system", "content": "Bạn là một chuyên gia biên tập kịch bản, luôn xuất sắc trong nhiệm vụ xây dựng kịch bản storyboard."},
                {"role": "user", "content": f"Hãy phân tích và chia đoạn văn bản sau thành các đoạn riêng biệt và tổng số đoạn không quá 6:\n\n{script_content}"}
            ],
            temperature=0.7,
            max_tokens=2000
        )
        paragraphs = response.choices[0].message.content.split("\n\n")
        paragraphs = [p.strip() for p in paragraphs if p.strip()]
        return paragraphs
    except Exception as e:
        return f"Lỗi khi phân tích và chia đoạn: {str(e)}"

def create_storyboard(script_content):
    # Phân chia script thành các đoạn
    paragraphs = split_into_paragraphs(script_content)
    
    # Tạo mô tả ảnh minh họa cho mỗi đoạn
    storyboard = []
    for paragraph in paragraphs:
        image_description = generate_image_description(paragraph)
        storyboard.append({
            "paragraph": paragraph,
            "image_description": image_description
        })
    
    return storyboard

def generate_storyboard(script_content):
    try:
        status = "Đang xử lý..."
        storyboard = create_storyboard(script_content)
        images = generate_images(storyboard)
        status = "Đã tạo Storyboard thành công!"
        
        # Chuyển đổi các đối tượng dict thành các đối tượng hình ảnh và văn bản tương ứng
        gallery_items = []
        for item in storyboard:
            paragraph = item["paragraph"]
            image_description = item["image_description"]
            gallery_items.append((images.pop(0), paragraph))
        
        return gallery_items, status
    except Exception as e:
        status = f"Đã xảy ra lỗi: {str(e)}"
        return [], status

def generate_image_description(paragraph):
    # Sử dụng LLM để tạo mô tả ảnh từ đoạn văn
    try:
        client = Groq(api_key=GROQ_API_KEY)
        response = client.chat.completions.create(
            model="mixtral-8x7b-32768",
            messages=[
                {"role": "system", "content": "Bạn là một chuyên gia sáng tạo nội dung, đặc biệt giỏi trong việc đưa ra ý tưởng về hình ảnh minh họa cho các đoạn văn hay kịch bản điện ảnh."},
                {"role": "user", "content": f"Tạo mô tả ảnh cho đoạn văn sau:\n\n{paragraph}"}
            ],
            temperature=0.7,
            max_tokens=200
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"Lỗi khi tạo mô tả ảnh: {str(e)}"

def query_image_generation(prompt):
    # Sử dụng DiffusionPipeline để tạo ảnh
    image = pipe(prompt).images[0]
    return image

def generate_images(storyboard):
    images = []
    for item in storyboard:
        image_description = item["image_description"]
        image = query_image_generation(image_description)
        images.append(image)
    return images

def create_video_clip(image, audio_path):
    image_clip = ImageClip(image)
    audio_clip = AudioFileClip(audio_path)
    video_clip = image_clip.set_audio(audio_clip)
    return video_clip

def create_final_video(storyboard, images, audio_paths):
    video_clips = []
    for i, item in enumerate(storyboard):
        paragraph = item["paragraph"]
        image = images[i]
        audio_path = audio_paths[i]
        video_clip = create_video_clip(image, audio_path)
        video_clips.append(video_clip)
    
    final_video = concatenate_videoclips(video_clips, method="compose")
    final_video_path = "final_video.mp4"
    final_video.write_videofile(final_video_path, codec="libx264")
    return final_video_path

def interface():
    with gr.Blocks() as app:
        gr.Markdown("# Ứng dụng Tạo Nội dung và Âm thanh")

        with gr.Tab("Tạo Nội dung"):
            prompt = gr.Textbox(label="Nhập yêu cầu nội dung")
            file_upload = gr.File(label="Tải lên file kèm theo", type="filepath")

            content_type = gr.Radio(
                label="Chọn loại nội dung",
                choices=CONTENT_TYPES,
                value=None
            )

            content_button = gr.Button("Tạo Nội dung")
            content_output = gr.Textbox(label="Nội dung tạo ra", interactive=True)
            confirm_button = gr.Button("Xác nhận nội dung")
            download_docx = gr.File(label="Tải xuống file DOCX", interactive=False)
            status_message = gr.Label(label="Trạng thái")

            def generate_content(prompt, file, content_type):
                try:
                    status = "Đang xử lý..."
                    if file and os.path.exists(file):
                        mime_type, _ = mimetypes.guess_type(file)
                        if mime_type == "application/pdf":
                            file_content = process_pdf(file)
                            prompt = f"{prompt}\n\nDưới đây là nội dung của file tài liệu:\n\n{file_content}"
                        elif mime_type in ("application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                           "application/msword"):
                            file_content = process_docx(file)
                            prompt = f"{prompt}\n\nDưới đây là nội dung của file tài liệu:\n\n{file_content}"
                        else:
                            raise ValueError("Định dạng file không được hỗ trợ.")

                    if not content_type:
                        raise ValueError("Vui lòng chọn một loại nội dung")

                    script_content = create_content(prompt, content_type, "Tiếng Việt")
                    docx_path = "script.docx"
                    create_docx(script_content, docx_path)

                    status = "Đã tạo nội dung thành công!"
                    return script_content, docx_path, status
                except Exception as e:
                    status = f"Đã xảy ra lỗi: {str(e)}"
                    return "", None, status

            async def confirm_content(content):
                docx_path = "script.docx"
                create_docx(content, docx_path)

                status = "Nội dung đã được xác nhận!"
                return docx_path, status

            content_button.click(generate_content,
                                 inputs=[prompt, file_upload, content_type],
                                 outputs=[content_output, download_docx, status_message])

            confirm_button.click(lambda x: asyncio.run(confirm_content(x)),
                                 inputs=[content_output],
                                 outputs=[download_docx, status_message])

        with gr.Tab("Storyboard"):
            gr.Markdown("## Tạo Storyboard từ Script")

            script_input = gr.Textbox(label="Nhập nội dung script", lines=10)
            storyboard_button = gr.Button("Tạo Storyboard")
            storyboard_output = gr.Gallery(label="Storyboard", columns=2, height="auto")
            confirm_storyboard_button = gr.Button("Xác nhận Storyboard")
            status_message_storyboard = gr.Label(label="Trạng thái")

            def generate_storyboard(script_content):
                try:
                    status = "Đang xử lý..."
                    storyboard = create_storyboard(script_content)
                    images = generate_images(storyboard)
                    status = "Đã tạo Storyboard thành công!"
                    
                    gallery_items = []
                    for item in storyboard:
                        paragraph = item["paragraph"]
                        image_description = item["image_description"]
                        gallery_items.append((images.pop(0), paragraph))
                    
                    return gallery_items, status
                except Exception as e:
                    status = f"Đã xảy ra lỗi: {str(e)}"
                    return [], status

            storyboard_button.click(generate_storyboard,
                                    inputs=[script_input],
                                    outputs=[storyboard_output, status_message_storyboard])

    return app

# Khởi chạy ứng dụng
if __name__ == "__main__":
    app = interface()
    app.launch()
import gradio as gr
import requests
import json
import os
import io
import tempfile
import mimetypes
import time
import asyncio
import functools
import shutil
import random
import glob
from typing import List

from groq import Groq
from docx import Document
import fitz  # PyMuPDF để xử lý PDF
from PIL import Image, ImageDraw, ImageFont
from sentence_transformers import SentenceTransformer, util
from transformers import AutoTokenizer, AutoModel
from moviepy.editor import ImageClip, AudioFileClip, concatenate_videoclips
from edge_tts import Communicate  # Sử dụng Edge TTS
from huggingface_hub import login
from diffusers import DiffusionPipeline

# Cấu hình API keys
GROQ_API_KEY = os.environ.get("GROQ_API_KEY")
HF_API_KEY = os.environ.get("HF_API_KEY")

# Danh sách giọng đọc và ngôn ngữ
VOICES = ["vi-VN-HoaiMyNeural"]  # Thay đổi giọng đọc Tiếng Việt
LANGUAGES = ["Tiếng Anh", "Tiếng Việt"]
CONTENT_TYPES = ["podcast", "giới thiệu", "triết lý sống", "Phổ biến kiến thức thống kê"]

# Hướng dẫn mặc định cho từng loại nội dung
CONTENT_TYPE_INSTRUCTIONS = {
    "podcast": """
    Tone giọng: Gần gũi, thân thiện nhưng chuyên sâu, thể hiện sự am hiểu về chủ đề.
    Cấu trúc: 
    - Bắt đầu bằng một câu hỏi kích thích tư duy hoặc một câu chuyện mở màn gây tò mò.
    - Triển khai các luận điểm theo từng bước. Sử dụng câu từ mạnh mẽ, ví dụ điển hình hoặc những câu nói nổi tiếng.
    - Xây dựng các phần chuyển tiếp mượt mà giữa các ý.
    - Kết thúc podcast với một thông điệp sâu sắc, để lại sự suy ngẫm cho thính giả.
    Mục tiêu: Mang lại kiến thức giá trị, lôi cuốn thính giả tham gia suy nghĩ và cảm nhận sâu sắc về chủ đề.
    """,
    "giới thiệu": """
    Tone giọng: Chuyên nghiệp, gãy gọn nhưng vẫn có sự truyền cảm.
    Cấu trúc:
    - Bắt đầu với một câu khẳng định mạnh mẽ về đối tượng được giới thiệu.
    - Giải thích mục tiêu của phần giới thiệu, nhấn mạnh tầm quan trọng hoặc sự khác biệt.
    - Kết thúc với một lời kêu gọi hành động, khích lệ người nghe tiếp tục lắng nghe hoặc tham gia.
    Mục tiêu: Đưa ra thông tin cô đọng, hấp dẫn, khiến người nghe cảm thấy bị thu hút và muốn tìm hiểu thêm.
    """,
    "triết lý sống": """
    Tone giọng: Sâu sắc, truyền cảm hứng, mang tính chiêm nghiệm.
    Cấu trúc:
    - Bắt đầu bằng một câu hỏi sâu sắc hoặc ẩn dụ về cuộc sống.
    - Triển khai các luận điểm chặt chẽ, xen lẫn cảm xúc và những ví dụ đời thực hoặc những câu nói triết lý.
    - Kết thúc với một thông điệp sâu sắc, khơi dậy suy ngẫm cho người nghe.
    Mục tiêu: Khơi gợi suy nghĩ sâu sắc về cuộc sống, khiến người nghe tìm thấy ý nghĩa hoặc giá trị trong câu chuyện.
    """,
    "Phổ biến kiến thức Thống kê": """
    Tone giọng: Thân thiện, dễ hiểu, và mang tính giáo dục.
    Cấu trúc:
    - Bắt đầu với một câu hỏi hoặc một tình huống thực tế để thu hút sự chú ý.
    - Giải thích các khái niệm thống kê cơ bản một cách đơn giản và dễ hiểu, sử dụng ví dụ thực tế để minh họa.
    - Đưa ra các ứng dụng thực tế của thống kê trong đời sống hàng ngày hoặc trong các lĩnh vực cụ thể.
    - Kết thúc với một thông điệp khuyến khích người nghe áp dụng kiến thức thống kê vào cuộc sống.
    Mục tiêu: Giúp người nghe hiểu và yêu thích thống kê, thấy được giá trị và ứng dụng của nó trong cuộc sống.
    """
}

# Khởi tạo model sentence transformer
model = SentenceTransformer('all-MiniLM-L6-v2')

# Đăng nhập vào Hugging Face Hub
login(token=HF_API_KEY)

# Khởi tạo DiffusionPipeline
pipe = DiffusionPipeline.from_pretrained("black-forest-labs/FLUX.1-dev")
pipe.load_lora_weights("TDN-M/RetouchFLux")

@functools.lru_cache(maxsize=32)
def create_content(prompt, content_type, language):
    content_type_instructions = CONTENT_TYPE_INSTRUCTIONS.get(content_type, "")
    general_instructions = f"""
    Viết một kịch bản dựa trên các ý chính và ý tưởng sáng tạo từ yêu cầu của người dùng. Sử dụng giọng điệu trò chuyện và bao gồm bất kỳ bối cảnh hoặc giải thích cần thiết nào để làm cho nội dung dễ tiếp cận với khán giả.
    Bắt đầu kịch bản bằng cách nêu rõ chủ đề, tham chiếu đến tiêu đề hoặc đề mục trong văn bản đầu vào. Nếu văn bản đầu vào không có tiêu đề, hãy đưa ra một tóm tắt ngắn gọn về nội dung được đề cập để mở đầu.
    Bao gồm các định nghĩa và thuật ngữ rõ ràng, cùng với ví dụ cho tất cả các vấn đề chính.
    Không bao gồm bất kỳ placeholder nào trong ngoặc vuông như [Host] hoặc [Guest]. Thiết kế đầu ra của bạn để được đọc to - nó sẽ được chuyển đổi trực tiếp thành âm thanh.
    Chỉ có một người nói, bạn. Giữ đúng chủ đề và duy trì một luồng hấp dẫn. Văn phong triển khai một cách tự nhiên, .
    Kịch bản nên có khoảng 1000 từ, chia thành 5 đoạn rõ ràng. Hãy tuân theo những hướng dẫn cụ thể sau cho thể loại {content_type}:
    {content_type_instructions}
    Ngôn ngữ sử dụng: {language}
    """

    try:
        client = Groq(api_key=GROQ_API_KEY)
        retries = 3
        for attempt in range(retries):
            try:
                chat_completion = client.chat.completions.create(
                    model="mixtral-8x7b-32768",
                    messages=[
                        {"role": "system", "content": general_instructions},
                        {"role": "user", "content": prompt}
                    ],
                    temperature=0.7,
                    max_tokens=8000
                )
                return chat_completion.choices[0].message.content
            except Exception as e:
                if attempt < retries - 1:
                    time.sleep(2 ** attempt)  # Exponential backoff
                else:
                    return f"Lỗi khi tạo nội dung: {str(e)}"
    except Exception as e:
        return f"Lỗi khi tạo nội dung: {str(e)}"

def text_to_speech(text, voice, language):
    try:
        communicate = Communicate(text, voice)
        with tempfile.NamedTemporaryFile(delete=False, suffix=".mp3") as temp_audio_file:
            communicate.save(temp_audio_file.name)
        return temp_audio_file.name
    except Exception as e:
        return f"Lỗi khi chuyển đổi văn bản thành giọng nói: {str(e)}"

def create_docx(content, output_path):
    doc = Document()
    doc.add_paragraph(content)
    doc.save(output_path)

def process_pdf(file_path):
    doc = fitz.open(file_path)
    text = ""
    for page in doc:
        text += page.get_text()
    return text

def process_docx(file_path):
    doc = Document(file_path)
    text = ""
    for para in doc.paragraphs:
        text += para.text
    return text

def split_into_paragraphs(script_content):
    try:
        client = Groq(api_key=GROQ_API_KEY)
        response = client.chat.completions.create(
            model="mixtral-8x7b-32768",
            messages=[
                {"role": "system", "content": "Bạn là một chuyên gia biên tập kịch bản, luôn xuất sắc trong nhiệm vụ xây dựng kịch bản storyboard."},
                {"role": "user", "content": f"Hãy phân tích và chia đoạn văn bản sau thành các đoạn riêng biệt và tổng số đoạn không quá 6:\n\n{script_content}"}
            ],
            temperature=0.7,
            max_tokens=2000
        )
        paragraphs = response.choices[0].message.content.split("\n\n")
        paragraphs = [p.strip() for p in paragraphs if p.strip()]
        return paragraphs
    except Exception as e:
        return f"Lỗi khi phân tích và chia đoạn: {str(e)}"

def create_storyboard(script_content):
    # Phân chia script thành các đoạn
    paragraphs = split_into_paragraphs(script_content)
    
    # Tạo mô tả ảnh minh họa cho mỗi đoạn
    storyboard = []
    for paragraph in paragraphs:
        image_description = generate_image_description(paragraph)
        storyboard.append({
            "paragraph": paragraph,
            "image_description": image_description
        })
    
    return storyboard

def generate_storyboard(script_content):
    try:
        status = "Đang xử lý..."
        storyboard = create_storyboard(script_content)
        images = generate_images(storyboard)
        status = "Đã tạo Storyboard thành công!"
        
        # Chuyển đổi các đối tượng dict thành các đối tượng hình ảnh và văn bản tương ứng
        gallery_items = []
        for item in storyboard:
            paragraph = item["paragraph"]
            image_description = item["image_description"]
            gallery_items.append((images.pop(0), paragraph))
        
        return gallery_items, status
    except Exception as e:
        status = f"Đã xảy ra lỗi: {str(e)}"
        return [], status

def generate_image_description(paragraph):
    # Sử dụng LLM để tạo mô tả ảnh từ đoạn văn
    try:
        client = Groq(api_key=GROQ_API_KEY)
        response = client.chat.completions.create(
            model="mixtral-8x7b-32768",
            messages=[
                {"role": "system", "content": "Bạn là một chuyên gia sáng tạo nội dung, đặc biệt giỏi trong việc đưa ra ý tưởng về hình ảnh minh họa cho các đoạn văn hay kịch bản điện ảnh."},
                {"role": "user", "content": f"Tạo mô tả ảnh cho đoạn văn sau:\n\n{paragraph}"}
            ],
            temperature=0.7,
            max_tokens=200
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"Lỗi khi tạo mô tả ảnh: {str(e)}"

def query_image_generation(prompt):
    # Sử dụng DiffusionPipeline để tạo ảnh
    image = pipe(prompt).images[0]
    return image

def generate_images(storyboard):
    images = []
    for item in storyboard:
        image_description = item["image_description"]
        image = query_image_generation(image_description)
        images.append(image)
    return images

def create_video_clip(image, audio_path):
    image_clip = ImageClip(image)
    audio_clip = AudioFileClip(audio_path)
    video_clip = image_clip.set_audio(audio_clip)
    return video_clip

def create_final_video(storyboard, images, audio_paths):
    video_clips = []
    for i, item in enumerate(storyboard):
        paragraph = item["paragraph"]
        image = images[i]
        audio_path = audio_paths[i]
        video_clip = create_video_clip(image, audio_path)
        video_clips.append(video_clip)
    
    final_video = concatenate_videoclips(video_clips, method="compose")
    final_video_path = "final_video.mp4"
    final_video.write_videofile(final_video_path, codec="libx264")
    return final_video_path

def interface():
    with gr.Blocks() as app:
        gr.Markdown("# Ứng dụng Tạo Nội dung và Âm thanh")

        with gr.Tab("Tạo Nội dung"):
            prompt = gr.Textbox(label="Nhập yêu cầu nội dung")
            file_upload = gr.File(label="Tải lên file kèm theo", type="filepath")

            content_type = gr.Radio(
                label="Chọn loại nội dung",
                choices=CONTENT_TYPES,
                value=None
            )

            content_button = gr.Button("Tạo Nội dung")
            content_output = gr.Textbox(label="Nội dung tạo ra", interactive=True)
            confirm_button = gr.Button("Xác nhận nội dung")
            download_docx = gr.File(label="Tải xuống file DOCX", interactive=False)
            status_message = gr.Label(label="Trạng thái")

            def generate_content(prompt, file, content_type):
                try:
                    status = "Đang xử lý..."
                    if file and os.path.exists(file):
                        mime_type, _ = mimetypes.guess_type(file)
                        if mime_type == "application/pdf":
                            file_content = process_pdf(file)
                            prompt = f"{prompt}\n\nDưới đây là nội dung của file tài liệu:\n\n{file_content}"
                        elif mime_type in ("application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                           "application/msword"):
                            file_content = process_docx(file)
                            prompt = f"{prompt}\n\nDưới đây là nội dung của file tài liệu:\n\n{file_content}"
                        else:
                            raise ValueError("Định dạng file không được hỗ trợ.")

                    if not content_type:
                        raise ValueError("Vui lòng chọn một loại nội dung")

                    script_content = create_content(prompt, content_type, "Tiếng Việt")
                    docx_path = "script.docx"
                    create_docx(script_content, docx_path)

                    status = "Đã tạo nội dung thành công!"
                    return script_content, docx_path, status
                except Exception as e:
                    status = f"Đã xảy ra lỗi: {str(e)}"
                    return "", None, status

            async def confirm_content(content):
                docx_path = "script.docx"
                create_docx(content, docx_path)

                status = "Nội dung đã được xác nhận!"
                return docx_path, status

            content_button.click(generate_content,
                                 inputs=[prompt, file_upload, content_type],
                                 outputs=[content_output, download_docx, status_message])

            confirm_button.click(lambda x: asyncio.run(confirm_content(x)),
                                 inputs=[content_output],
                                 outputs=[download_docx, status_message])

        with gr.Tab("Storyboard"):
            gr.Markdown("## Tạo Storyboard từ Script")

            script_input = gr.Textbox(label="Nhập nội dung script", lines=10)
            storyboard_button = gr.Button("Tạo Storyboard")
            storyboard_output = gr.Gallery(label="Storyboard", columns=2, height="auto")
            confirm_storyboard_button = gr.Button("Xác nhận Storyboard")
            status_message_storyboard = gr.Label(label="Trạng thái")

            def generate_storyboard(script_content):
                try:
                    status = "Đang xử lý..."
                    storyboard = create_storyboard(script_content)
                    images = generate_images(storyboard)
                    status = "Đã tạo Storyboard thành công!"
                    
                    gallery_items = []
                    for item in storyboard:
                        paragraph = item["paragraph"]
                        image_description = item["image_description"]
                        gallery_items.append((images.pop(0), paragraph))
                    
                    return gallery_items, status
                except Exception as e:
                    status = f"Đã xảy ra lỗi: {str(e)}"
                    return [], status

            storyboard_button.click(generate_storyboard,
                                    inputs=[script_input],
                                    outputs=[storyboard_output, status_message_storyboard])

    return app

# Khởi chạy ứng dụng
if __name__ == "__main__":
    app = interface()
    app.launch()
